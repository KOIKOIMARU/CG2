from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "generated" / "assignments"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "LE3B_07_コイズミ_リョウ_追加要素案出し.docx"
PDF_PATH = OUT_DIR / "LE3B_07_コイズミ_リョウ_追加要素案出し.pdf"
FONT_PATH = Path(r"C:\Windows\Fonts\NotoSansJP-VF.ttf")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_table_layout(table, width_cm):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")


def style_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(cell, text, size=9.0, bold=False):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    p.text = ""
    run = p.add_run(text)
    style_run(run, size=size, bold=bold)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    style_run(run, size=16 if level == 1 else 13, bold=True, color="1F4E79")
    return p


ideas = [
    ("操作", "ロックオン照準", "照準を敵に近づけると一定範囲内の敵を自動で捕捉し、ロックオン状態で攻撃できるようにする。", "狙う気持ちよさを上げる。優先度:高 / 難易度:中"),
    ("操作", "チャージショット強化", "長押し時間に応じて弾のサイズ、威力、貫通数、エフェクトを段階的に変化させる。", "主力攻撃にメリハリを作る。優先度:高 / 難易度:中"),
    ("操作", "回避アクション", "短い無敵時間のある左右ステップやバレルロールを追加する。", "敵弾を避ける遊びを増やす。優先度:高 / 難易度:中"),
    ("操作", "照準慣性", "マウス照準に少しだけ追従遅れを入れ、急操作時も滑らかに動くようにする。", "空中戦らしい操作感を出す。優先度:中 / 難易度:低"),
    ("操作", "オーバーヒート", "連射し続けると武器が一時的に撃てなくなり、撃つ間隔の管理が必要になる。", "連射一辺倒を防ぐ。優先度:中 / 難易度:低"),
    ("操作", "近距離迎撃ショット", "敵が接近したときだけ強い短射程攻撃を出せるようにする。", "接近された時の対応手段を作る。優先度:中 / 難易度:中"),
    ("操作", "スコア倍率攻撃", "連続ヒット中に特定タイミングで撃つとスコア倍率が上がる。", "上達要素を作る。優先度:中 / 難易度:中"),
    ("操作", "自動連射と手動射撃の切替", "通常時は自動連射、クリック長押しでチャージなど、攻撃操作を整理する。", "遊びやすさを上げる。優先度:高 / 難易度:低"),
    ("敵AI", "蛇行移動する敵", "奥からまっすぐ来るだけでなく、左右に曲線移動しながら接近する敵を作る。", "単調さを減らす。優先度:高 / 難易度:低"),
    ("敵AI", "編隊移動", "複数の敵がV字、横一列、円形などの隊列を組んで動く。", "画面映えと攻略感を出す。優先度:高 / 難易度:中"),
    ("敵AI", "一撃離脱型の敵", "接近して攻撃した後、画面外へ離脱する敵を作る。", "敵に個性を出す。優先度:高 / 難易度:中"),
    ("敵AI", "遠距離射撃型の敵", "プレイヤーに接近せず、一定距離を保って弾を撃つ。", "敵の役割分担を作る。優先度:高 / 難易度:中"),
    ("敵AI", "シールド持ちの敵", "正面からの攻撃を防ぎ、横移動中や攻撃直後だけ弱点が出る。", "狙う場所を考えさせる。優先度:中 / 難易度:中"),
    ("敵AI", "分裂する敵", "倒すと小さい敵に分裂し、追撃が必要になる。", "撃破後の展開を増やす。優先度:中 / 難易度:中"),
    ("敵AI", "突進予告", "突進前に光や音で予兆を出し、その後高速で接近する。", "理不尽さを減らして緊張感を出す。優先度:高 / 難易度:中"),
    ("敵AI", "弱点部位", "敵の一部だけに高ダメージ判定を持たせる。", "照準操作の意味を強くする。優先度:中 / 難易度:中"),
    ("敵AI", "敵ごとのHP差", "雑魚、硬い敵、速い敵などで耐久値を変える。", "敵バリエーションの基礎になる。優先度:高 / 難易度:低"),
    ("敵AI", "攻撃パターン切替", "HPが減る、時間が経つ、距離が近づくなどの条件で移動や攻撃を変える。", "敵を賢く見せる。優先度:高 / 難易度:中"),
    ("ウェーブ", "ウェーブごとのテーマ", "1波目は雑魚、2波目は編隊、3波目は突進型など、波ごとに役割を持たせる。", "ゲーム進行を分かりやすくする。優先度:高 / 難易度:低"),
    ("ウェーブ", "中ボスウェーブ", "数ウェーブごとに大型の敵を出し、攻撃パターンを変化させる。", "区切りと達成感を作る。優先度:高 / 難易度:高"),
    ("ウェーブ", "時間制限ミッション", "一定時間以内に敵を倒すとボーナスが入る。", "テンポを上げる。優先度:中 / 難易度:低"),
    ("ウェーブ", "護衛対象", "画面内の味方やオブジェクトを守りながら戦うウェーブを作る。", "遊びの種類を増やす。優先度:低 / 難易度:高"),
    ("ウェーブ", "分岐ルート", "撃破数やスコアに応じて次の敵構成や背景が変化する。", "周回性を作る。優先度:中 / 難易度:高"),
    ("ウェーブ", "難易度自動調整", "被弾が多い時は敵数を減らし、上手い時は敵を増やす。", "遊びやすさを保つ。優先度:中 / 難易度:中"),
    ("演出", "ヒットストップ", "強い攻撃が当たった瞬間に一瞬だけ時間を止める。", "攻撃の手応えを強くする。優先度:高 / 難易度:低"),
    ("演出", "画面揺れ", "被弾、爆発、チャージショット命中時にカメラを短く揺らす。", "迫力を出す。優先度:高 / 難易度:低"),
    ("演出", "敵撃破エフェクト強化", "敵の種類ごとに爆発色、粒子、残像を変える。", "倒した気持ちよさを上げる。優先度:高 / 難易度:中"),
    ("演出", "低HP演出", "HPが少ない時に画面端を赤くし、音やUIで危険状態を伝える。", "状況を分かりやすくする。優先度:高 / 難易度:低"),
    ("演出", "スローモーション演出", "ウェーブクリアや強敵撃破時に一瞬だけスローにする。", "見せ場を作る。優先度:中 / 難易度:中"),
    ("演出", "レール移動カメラ演出", "場面に応じてカメラの高さや角度を少し変化させる。", "空中戦の疾走感を出す。優先度:高 / 難易度:中"),
    ("UI", "敵ロックオン表示", "ロックオン中の敵に枠やマーカーを表示する。", "狙っている対象を明確にする。優先度:高 / 難易度:低"),
    ("UI", "ウェーブ表示", "現在のウェーブ数、残り敵数、次ウェーブまでの時間を表示する。", "進行状況を伝える。優先度:高 / 難易度:低"),
    ("UI", "チャージゲージ", "チャージショットの溜まり具合を照準付近や画面下に表示する。", "操作状態を分かりやすくする。優先度:高 / 難易度:低"),
    ("UI", "スコア表示", "撃破、連続ヒット、ノーダメージなどでスコアを加算して表示する。", "遊びの目標を増やす。優先度:中 / 難易度:低"),
    ("UI", "コンボ表示", "短時間に敵を倒し続けるとコンボ数と倍率を表示する。", "上手く遊ぶ動機を作る。優先度:中 / 難易度:中"),
    ("UI", "被弾方向インジケータ", "どの方向から攻撃を受けたか画面端に表示する。", "状況把握を助ける。優先度:中 / 難易度:中"),
    ("成長", "武器レベルアップ", "敵撃破やアイテム取得で一定時間だけ弾速、威力、連射力が上がる。", "強くなる楽しさを作る。優先度:中 / 難易度:中"),
    ("成長", "一時強化アイテム", "回復、攻撃力上昇、シールド、スローなどのアイテムを出現させる。", "戦闘中の変化を増やす。優先度:中 / 難易度:中"),
    ("成長", "リザルト評価", "クリア時に命中率、被弾数、撃破数、スコアでランクを出す。", "リプレイ性を作る。優先度:高 / 難易度:中"),
    ("成長", "実績条件", "ノーダメージ、全敵撃破、チャージのみクリアなどの条件を記録する。", "やり込み要素を増やす。優先度:低 / 難易度:中"),
    ("ステージ", "背景オブジェクト通過", "雲、柱、岩、リングなどを通過させて前進感を強める。", "レールシューティングらしさを上げる。優先度:高 / 難易度:中"),
    ("ステージ", "障害物回避", "プレイヤーが左右移動で避ける障害物を配置する。", "撃つ以外の操作を増やす。優先度:中 / 難易度:中"),
    ("ステージ", "チェックポイント", "一定区間ごとに復帰地点を保存し、失敗時に途中から再開する。", "遊びやすさを上げる。優先度:中 / 難易度:中"),
    ("ステージ", "背景変化", "ウェーブ進行に合わせて空の色、雲量、光の向きを変える。", "進行感と雰囲気を出す。優先度:中 / 難易度:中"),
    ("サウンド", "ヒット音の種類分け", "通常命中、弱点命中、撃破、チャージ命中で効果音を変える。", "操作の手応えを音でも伝える。優先度:高 / 難易度:低"),
    ("サウンド", "危険時BGM変化", "HP低下や中ボス出現時にBGMや効果音を変える。", "場面の緊張感を作る。優先度:中 / 難易度:中"),
    ("システム", "デバッグ表示切替", "敵数、FPS、現在ウェーブ、プレイヤーHPなどを開発用に表示する。", "調整をしやすくする。優先度:高 / 難易度:低"),
    ("システム", "敵出現データの外部化", "敵の種類、出現時間、位置、移動パターンをJSONなどで管理する。", "調整と量産をしやすくする。優先度:高 / 難易度:中"),
    ("システム", "リプレイ保存", "入力や結果を記録し、あとで再生できるようにする。", "デバッグと作品紹介に使える。優先度:低 / 難易度:高"),
    ("システム", "チュートリアル", "照準、射撃、チャージ、回避を順番に説明する短い導入を作る。", "初見でも遊びやすくする。優先度:中 / 難易度:中"),
]


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.7)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Yu Gothic"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
normal.font.size = Pt(10.5)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(8)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("個人制作ゲームの追加要素案出し")
style_run(run, size=22, bold=True, color="000000")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("LE3B_07_コイズミ_リョウ")
style_run(run, size=11, bold=False, color="555555")

add_heading(doc, "1. 課題内容", 1)
p = doc.add_paragraph()
r = p.add_run(
    "自分が制作している個人制作ゲームについて、ChatGPTを用いて追加要素の案出しを行った。"
    "今回は、実装にプログラミングが必要な要素に限定し、ゲームとしての完成度、操作感、敵AI、演出、UIの強化につながる案を出してもらうことを目的とした。"
)
style_run(r)

add_heading(doc, "2. AIに渡したプロンプト", 1)
prompt = (
    "以下で説明するゲームに追加する要素を50個挙げてください。実装にプログラミングが必要なものに限定してください。\n\n"
    "現在、3Dレールシューティングゲームを制作しています。参考にしているのは「新・光神話 パルテナの鏡」の空中戦パートです。"
    "プレイヤーは自動で前方へ進むような構成で、マウス操作で照準を動かし、敵を撃ち落とします。"
    "現在実装済みの要素は、タイトル画面、ゲーム画面、プレイヤー、敵、弾、HPバー、ウェーブ制、ポストエフェクト、"
    "スカイボックス、簡単な敵AI、マウス照準、チャージショット、ヒット演出などです。\n\n"
    "ゲームとしては、単に敵が前から来るだけではなく、レールシューティングらしいテンポ、爽快感、敵の動きの変化、"
    "演出、成長要素などを増やして完成度を上げたいです。一人で制作しているため、実装の優先度や難易度も分かるようにしてください。"
    "各案は「追加要素名」「内容」「実装の狙い」の形で書いてください。"
)
box = doc.add_table(rows=1, cols=1)
box.alignment = WD_TABLE_ALIGNMENT.CENTER
set_fixed_table_layout(box, 17.6)
set_table_borders(box, color="D9E2F3", size="6")
cell = box.cell(0, 0)
set_cell_shading(cell, "F4F8FF")
set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
add_para(cell, prompt, size=9.4)

add_heading(doc, "3. ChatGPTによる追加要素案", 1)
p = doc.add_paragraph()
r = p.add_run("以下は、上記のプロンプトに対して得られた追加要素案である。")
style_run(r)

table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
set_fixed_table_layout(table, 17.6)
set_table_borders(table)
widths = [0.9, 1.8, 3.2, 7.1, 4.6]
headers = ["No.", "分類", "追加要素", "内容", "狙い・優先度"]
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_width(cell, widths[i])
    set_cell_shading(cell, "1F4E79")
    set_cell_margins(cell, top=90, bottom=90, start=80, end=80)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    style_run(run, size=8.8, bold=True, color="FFFFFF")

for index, (category, name, content, aim) in enumerate(ideas, start=1):
    row = table.add_row()
    values = [str(index), category, name, content, aim]
    for i, text in enumerate(values):
        cell = row.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_margins(cell, top=70, bottom=70, start=70, end=70)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if index % 2 == 0:
            set_cell_shading(cell, "F7F9FC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(text)
        style_run(run, size=7.8 if i in (3, 4) else 8.0, bold=(i == 2))

add_heading(doc, "4. 優先して実装したい要素", 1)
priority_items = [
    "ロックオン照準、チャージゲージ、敵ロックオン表示を組み合わせることで、プレイヤーの攻撃操作を分かりやすくする。",
    "蛇行移動、編隊移動、一撃離脱型、遠距離射撃型などを追加し、敵が全員同じ動きに見えないようにする。",
    "ヒットストップ、画面揺れ、撃破エフェクト、低HP演出を入れ、攻撃や被弾の手応えを強める。",
    "ウェーブごとのテーマや中ボスウェーブを作り、ゲーム進行に区切りと達成感を持たせる。",
]
for item in priority_items:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(-0.2)
    r = p.add_run("・" + item)
    style_run(r)

add_heading(doc, "5. まとめ", 1)
p = doc.add_paragraph()
r = p.add_run(
    "AIを使用することで、短時間で多くの追加要素案を得ることができた。"
    "ただし、すべての案をそのまま使うのではなく、自分のゲームの方向性や制作期間、実装難易度に合わせて取捨選択する必要がある。"
    "今回の案の中では、敵AIのバリエーション、ロックオン照準、ウェーブ演出、ヒット演出の強化が特に現在のゲームに合っていると考えた。"
)
style_run(r)

doc.save(DOCX_PATH)


def build_pdf():
    font_name = "NotoSansJP"
    pdfmetrics.registerFont(TTFont(font_name, str(FONT_PATH)))

    pdf_doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=landscape(A4),
        rightMargin=1.1 * cm,
        leftMargin=1.1 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm,
        title="個人制作ゲームの追加要素案出し",
        author="LE3B_07_コイズミ_リョウ",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "JapaneseTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=20,
        leading=25,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        "JapaneseMeta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9.5,
        leading=13,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"),
        spaceAfter=14,
    )
    heading_style = ParagraphStyle(
        "JapaneseHeading",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#1F4E79"),
        spaceBefore=10,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "JapaneseBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=6,
    )
    small_style = ParagraphStyle(
        "JapaneseSmall",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=7.0,
        leading=9.5,
        alignment=TA_LEFT,
        wordWrap="CJK",
    )
    small_center_style = ParagraphStyle(
        "JapaneseSmallCenter",
        parent=small_style,
        alignment=TA_CENTER,
    )
    header_style = ParagraphStyle(
        "JapaneseTableHeader",
        parent=small_center_style,
        fontName=font_name,
        fontSize=7.4,
        leading=9,
        textColor=colors.white,
    )

    story = []
    story.append(Paragraph("個人制作ゲームの追加要素案出し", title_style))
    story.append(Paragraph("LE3B_07_コイズミ_リョウ", meta_style))

    story.append(Paragraph("1. 課題内容", heading_style))
    story.append(Paragraph(
        "自分が制作している個人制作ゲームについて、ChatGPTを用いて追加要素の案出しを行った。"
        "今回は、実装にプログラミングが必要な要素に限定し、ゲームとしての完成度、操作感、敵AI、演出、UIの強化につながる案を出してもらうことを目的とした。",
        body_style,
    ))

    story.append(Paragraph("2. AIに渡したプロンプト", heading_style))
    prompt_table = Table(
        [[Paragraph(prompt.replace("\n", "<br/>"), body_style)]],
        colWidths=[26.2 * cm],
    )
    prompt_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F8FF")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#D9E2F3")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(prompt_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph("3. ChatGPTによる追加要素案", heading_style))
    story.append(Paragraph("以下は、上記のプロンプトに対して得られた追加要素案である。", body_style))

    data = [[
        Paragraph("No.", header_style),
        Paragraph("分類", header_style),
        Paragraph("追加要素", header_style),
        Paragraph("内容", header_style),
        Paragraph("狙い・優先度", header_style),
    ]]
    for index, (category, name, content, aim) in enumerate(ideas, start=1):
        data.append([
            Paragraph(str(index), small_center_style),
            Paragraph(category, small_center_style),
            Paragraph(name, small_style),
            Paragraph(content, small_style),
            Paragraph(aim, small_style),
        ])

    ideas_table = Table(
        data,
        colWidths=[0.8 * cm, 1.6 * cm, 3.2 * cm, 11.5 * cm, 9.1 * cm],
        repeatRows=1,
        splitByRow=1,
    )
    table_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    for row_index in range(1, len(data)):
        if row_index % 2 == 0:
            table_style.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor("#F7F9FC")))
    ideas_table.setStyle(TableStyle(table_style))
    story.append(ideas_table)

    story.append(PageBreak())
    story.append(Paragraph("4. 優先して実装したい要素", heading_style))
    priority_rows = []
    for item in priority_items:
        priority_rows.append([Paragraph("・", body_style), Paragraph(item, body_style)])
    priority_table = Table(priority_rows, colWidths=[0.5 * cm, 25.7 * cm])
    priority_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
    ]))
    story.append(priority_table)

    story.append(Paragraph("5. まとめ", heading_style))
    story.append(Paragraph(
        "AIを使用することで、短時間で多くの追加要素案を得ることができた。"
        "ただし、すべての案をそのまま使うのではなく、自分のゲームの方向性や制作期間、実装難易度に合わせて取捨選択する必要がある。"
        "今回の案の中では、敵AIのバリエーション、ロックオン照準、ウェーブ演出、ヒット演出の強化が特に現在のゲームに合っていると考えた。",
        body_style,
    ))

    def add_page_number(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawRightString(28.6 * cm, 0.55 * cm, f"{doc_obj.page}")
        canvas.restoreState()

    pdf_doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


build_pdf()
print(DOCX_PATH)
print(PDF_PATH)
